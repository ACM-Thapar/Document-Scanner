import cv2
import numpy as np
import os
import pytesseract
import docx
import RegionSelector

pytesseract.pytesseract.tesseract_cmd = 'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'

path = 'TargetFile'
myPic = os.listdir(path)
imgQ = cv2.imread('TargetFile\\'+myPic[0])

per = 25

orb = cv2.ORB_create(1000)
kp1,des1 = orb.detectAndCompute(imgQ,None)


for j,y in enumerate(myPic):
    img = cv2.imread(path + "/" + y)
    #cv2.imshow(y,img)
    kp2,des2 = orb.detectAndCompute(img,None)
    bf = cv2.BFMatcher(cv2.NORM_HAMMING)
    matches = bf.match(des2,des1)
    matches = list(matches)
    matches.sort(key= lambda x: x.distance)
    good = matches[:int(len(matches)*(per/100))]
    imgMatch = cv2.drawMatches(img,kp2,imgQ,kp1,good[:100],None,flags=2)
    #cv2.imshow(y,imgMatch)

    srcPoints = np.float32([kp2[m.queryIdx].pt for m in good]).reshape(-1,1,2)
    dstPoints = np.float32([kp2[m.queryIdx].pt for m in good]).reshape(-1,1,2)

    M,_ = cv2.findHomography(srcPoints,dstPoints,cv2.RANSAC,5.0)
    h, w, c = img.shape
    imgScan = cv2.warpPerspective(img,M,(w,h))
    #cv2.imshow(y, imgScan)

    imgShow = imgScan.copy()
    imgMask = np.zeros_like(imgShow)

    myData = []

    print(f'#######Extracting Data from form {j}#######')

    roi = RegionSelector.function(y)

    for x,r in enumerate(roi):
        cv2.rectangle(imgMask,(r[0][0],r[0][1]),(r[1][0],r[1][1]),(0,255,0),cv2.FILLED)
        imgShow = cv2.addWeighted(imgShow,0.99,imgMask,0.1,0)

    imgCrop = imgScan[r[0][1]:r[1][1], r[0][0]:r[1][0]]
    cv2.imshow(str(x),imgCrop)

    print(f'{r[3]} :{pytesseract.image_to_string(imgCrop)}')
    myData.append(pytesseract.image_to_string(imgCrop))

mydoc = docx.Document()
for i in range(len(myData)):
    mydoc.add_paragraph(myData[i])

mydoc.save("my_written_file.docx")

'''doc = aw.Document()
builder = aw.DocumentBuilder(doc)

for i in range(len(myData)):
    builder.write(myData[i])

doc.save(myPicList[0]+".docx")'''

#cv2.imshow("KeyPointsQuery",imgKp1)
#cv2.imshow("Output",imgQ)
cv2.waitKey(0)

